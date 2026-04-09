
"""LogPeg runtime: auto-generated construction logs."""

from __future__ import annotations

import calendar
from collections import defaultdict
from datetime import UTC, date, datetime, timedelta
import hashlib
from io import BytesIO
import json
from typing import Any
from zoneinfo import ZoneInfo

from fastapi import HTTPException
import httpx
from reportlab.graphics import renderPDF
from reportlab.graphics.barcode import qr
from reportlab.graphics.shapes import Drawing
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from services.api.domain.logpeg.models import (
    AggregateSummary,
    Anomaly,
    CostSummary,
    DailyActivity,
    DailyLog,
    EquipmentSummary,
    LogPegAutoGenerateResult,
    MaterialConsumed,
    MaterialSummary,
    MonthlyLog,
    PersonnelSummary,
    ProgressSummary,
    QualitySummary,
    WeeklyLog,
)
from services.api.domain.utxo.integrations import ProofUTXOEngine

_SH_TZ = ZoneInfo("Asia/Shanghai")


def _txt(v: Any, d: str = "") -> str:
    if v is None:
        return d
    if isinstance(v, bytes):
        return v.decode("utf-8", errors="replace")
    return str(v)


def _num(v: Any, d: float = 0.0) -> float:
    try:
        return float(v)
    except Exception:
        return float(d)


def _dict(v: Any) -> dict[str, Any]:
    return v if isinstance(v, dict) else {}


def _list(v: Any) -> list[Any]:
    return v if isinstance(v, list) else []


def _round(v: float) -> float:
    return round(float(v or 0.0) + 1e-12, 2)


def _dt(v: Any) -> datetime:
    if isinstance(v, datetime):
        return v.astimezone(UTC) if v.tzinfo else v.replace(tzinfo=UTC)
    s = _txt(v).strip()
    if not s:
        return datetime.now(UTC)
    s = s[:-1] + "+00:00" if s.endswith("Z") else s
    d = datetime.fromisoformat(s)
    return d.astimezone(UTC) if d.tzinfo else d.replace(tzinfo=UTC)


def _local_date(v: Any) -> str:
    return _dt(v).astimezone(_SH_TZ).date().isoformat()


def _parse_date(v: str) -> date:
    try:
        return datetime.strptime(_txt(v).strip(), "%Y-%m-%d").date()
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=f"invalid_date: {v}") from exc


def _gate_text(value: str, language: str) -> str:
    key = _txt(value).strip().lower()
    if language == "en":
        if key in {"pass", "approved", "ok", "合格"}:
            return "Pass"
        if key in {"fail", "failed", "rejected", "不合格"}:
            return "Fail"
        return "Pending"
    if key in {"pass", "approved", "ok", "合格"}:
        return "合格"
    if key in {"fail", "failed", "rejected", "不合格"}:
        return "不合格"
    return "待定"


def _is_gate_fail(value: str) -> bool:
    return _txt(value).strip().lower() in {"fail", "failed", "rejected", "不合格"}


def _safe_file_token(value: str) -> str:
    token = _txt(value).strip().replace("/", "-").replace("\\", "-")
    return "".join(ch if ch.isalnum() or ch in {"-", "_", "."} else "-" for ch in token) or "log"


class LogPegEngine:
    def __init__(self, *, sb: Any) -> None:
        self.sb = sb

    def _project(self, project_uri: str) -> dict[str, Any]:
        rows = self.sb.table("projects").select("*").eq("v_uri", _txt(project_uri).strip().rstrip("/")).limit(1).execute().data or []
        if not rows:
            raise HTTPException(status_code=404, detail="project_not_found")
        return rows[0] if isinstance(rows[0], dict) else {}

    def _signed_row(self, project_uri: str, log_date: str) -> dict[str, Any] | None:
        rows = self.sb.table("proof_utxo").select("*").eq("project_uri", project_uri).eq("proof_type", "document").order("created_at", desc=True).limit(500).execute().data or []
        for row in rows:
            sd = _dict(row.get("state_data")) if isinstance(row, dict) else {}
            if _txt(sd.get("proof_kind")) == "logpeg_daily_signed" and _txt(sd.get("log_date") or sd.get("date")) == log_date:
                return row
        return None

    def _write_draft_once(self, *, project_uri: str, log: DailyLog) -> None:
        rows = (
            self.sb.table("proof_utxo")
            .select("proof_id,state_data")
            .eq("project_uri", project_uri)
            .eq("proof_type", "document")
            .order("created_at", desc=True)
            .limit(200)
            .execute()
            .data
            or []
        )
        for row in rows:
            if not isinstance(row, dict):
                continue
            sd = _dict(row.get("state_data"))
            if _txt(sd.get("proof_kind")) == "logpeg_daily_draft" and _txt(sd.get("log_date") or sd.get("date")) == log.log_date:
                return
        proof_id = f"GP-LOGPEG-DRAFT-{hashlib.sha256(f'{project_uri}:{log.log_date}:draft'.encode('utf-8')).hexdigest()[:16].upper()}"
        ProofUTXOEngine(self.sb).create(
            proof_id=proof_id,
            owner_uri=project_uri,
            project_uri=project_uri,
            proof_type="document",
            result="PENDING",
            state_data={"proof_kind": "logpeg_daily_draft", "log_date": log.log_date, "date": log.log_date, "log": log.model_dump(mode="json")},
            norm_uri="v://norm/NormPeg/LogPegDraft/1.0",
            segment_uri=f"{project_uri}/log/{log.log_date}",
            signer_uri=f"{project_uri}/role/system/",
            signer_role="SYSTEM",
        )

    def _log_hash(self, log: DailyLog) -> str:
        raw = json.dumps(log.model_dump(mode="json"), ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str)
        return hashlib.sha256(raw.encode("utf-8")).hexdigest()

    async def generate_daily_log(self, project_uri: str, log_date: str, weather: str = "", temperature_range: str = "", wind_level: str = "", language: str = "zh") -> DailyLog:
        d = _parse_date(log_date).isoformat()
        uri = _txt(project_uri).strip().rstrip("/")
        signed = self._signed_row(uri, d)
        if signed:
            sd = _dict(_dict(signed).get("state_data"))
            payload = _dict(sd.get("log"))
            if payload:
                out = DailyLog.model_validate(payload)
                out.locked = True
                out.sign_proof = _txt(_dict(signed).get("proof_id"))
                return out

        p = self._project(uri)
        trips = [
            row for row in (self.sb.table("gate_trips").select("*").order("signed_at", desc=False).limit(5000).execute().data or [])
            if isinstance(row, dict)
            and _local_date(row.get("signed_at") or row.get("created_at")) == d
            and (
                _txt(_dict(row.get("metadata")).get("project_uri")).strip().rstrip("/") == uri
                or _txt(row.get("trip_uri")).startswith(uri)
            )
        ]
        materials = [
            row for row in (self.sb.table("proof_utxo").select("*").eq("project_uri", uri).eq("proof_type", "inspection").order("created_at", desc=False).limit(5000).execute().data or [])
            if isinstance(row, dict)
            and _txt(_dict(row.get("state_data")).get("proof_kind")) in {"material_inspection_batch", "material_utxo_consume", "iqc_material_usage"}
            and _local_date(_dict(row.get("state_data")).get("created_at") or row.get("created_at")) == d
        ]
        nodes = [
            row for row in (self.sb.table("proof_utxo").select("*").eq("project_uri", uri).eq("proof_type", "node").order("created_at", desc=False).limit(10000).execute().data or [])
            if isinstance(row, dict) and _txt(_dict(row.get("state_data")).get("entity_type")) == "process_chain" and _dt(row.get("created_at")).astimezone(_SH_TZ).date() <= _parse_date(d)
        ]
        settles = [
            row for row in (self.sb.table("railpact_settlements").select("*").order("settled_at", desc=False).limit(20000).execute().data or [])
            if isinstance(row, dict) and (_txt(row.get("trip_uri")).startswith(uri) or _txt(row.get("project_uri")).strip().rstrip("/") == uri)
        ]
        exec_rows = [row for row in (self.sb.table("san_executors").select("*").limit(5000).execute().data or []) if isinstance(row, dict)]
        exec_map = {_txt(row.get("executor_uri")): row for row in exec_rows}
        trip_costs: dict[str, dict[str, float]] = defaultdict(lambda: {"labor": 0.0, "equipment": 0.0, "material": 0.0})
        daily_labor = daily_equipment = daily_material = cumulative_total = 0.0
        for row in settles:
            amount = _num(row.get("amount"))
            bucket = _txt(_dict(row.get("metadata")).get("smu_type") or row.get("smu_type") or "labor").lower()
            kind = "material" if "material" in bucket or "consumable" in bucket else "equipment" if any(x in bucket for x in ("equipment", "machine", "tool", "depreciation")) else "labor"
            when = _dt(row.get("settled_at") or row.get("created_at")).astimezone(_SH_TZ).date()
            if when <= _parse_date(d):
                cumulative_total += amount
            if when == _parse_date(d):
                if kind == "material":
                    daily_material += amount
                elif kind == "equipment":
                    daily_equipment += amount
                else:
                    daily_labor += amount
                t = _txt(row.get("trip_uri"))
                if t:
                    trip_costs[t][kind] += amount

        mat_by_component: dict[str, list[MaterialConsumed]] = defaultdict(list)
        mat_sum: dict[tuple[str, str, str], dict[str, float]] = defaultdict(lambda: {"qty": 0.0, "cost": 0.0})
        for row in materials:
            sd = _dict(row.get("state_data"))
            qty = _num(sd.get("quantity") or sd.get("actual_qty") or sd.get("used_qty"))
            unit_price = _num(sd.get("unit_price") or sd.get("price"))
            total = _round(_num(sd.get("cost"), qty * unit_price))
            comp = _txt(sd.get("component_uri"))
            rec = MaterialConsumed(name=_txt(sd.get("material_name") or sd.get("material_code")), code=_txt(sd.get("material_code")), qty=qty, unit=_txt(sd.get("unit")), unit_price=unit_price, total_cost=total, iqc_batch=_txt(sd.get("batch_no") or sd.get("inspection_batch_no")))
            mat_by_component[comp].append(rec)
            k = (rec.name, rec.code, rec.unit)
            mat_sum[k]["qty"] += rec.qty
            mat_sum[k]["cost"] += rec.total_cost

        latest_node: dict[str, dict[str, Any]] = {}
        for row in nodes:
            sd = _dict(row.get("state_data"))
            comp = _txt(sd.get("component_uri")).rstrip("/")
            if not comp:
                continue
            prev = latest_node.get(comp)
            if prev is None or _dt(row.get("created_at")) >= _dt(prev.get("created_at")):
                latest_node[comp] = row
        snapshot: dict[str, Any] = {}
        for comp, row in latest_node.items():
            sd = _dict(row.get("state_data"))
            sm = _dict(sd.get("state_matrix"))
            total_steps = int(sm.get("total_steps") or 0)
            completed_steps = int(sm.get("completed_steps") or 0)
            key = comp.rsplit("/", 1)[-1] if comp else comp
            snapshot[key] = {"component_uri": comp, "current_step": _txt(sd.get("current_step") or sm.get("current_step")), "completed_steps": completed_steps, "total_steps": total_steps, "status": "completed" if total_steps > 0 and completed_steps >= total_steps else "active"}

        activities: list[DailyActivity] = []
        for row in trips:
            md = _dict(row.get("metadata"))
            comp = _txt(md.get("component_uri"))
            trip_uri = _txt(row.get("trip_uri"))
            ex_uri = _txt(row.get("executor_uri") or md.get("executor_uri"))
            ex = _dict(exec_map.get(ex_uri))
            cost = trip_costs.get(trip_uri, {"labor": 0.0, "equipment": 0.0, "material": 0.0})
            gate = "不合格" if (_txt(row.get("action")).lower() == "reject" or not bool(row.get("verified", True))) else "合格"
            equips = md.get("equipment_used") if isinstance(md.get("equipment_used"), list) else []
            if isinstance(md.get("equipment_used"), str):
                equips = [x.strip() for x in _txt(md.get("equipment_used")).split(",") if x.strip()]
            if _txt(md.get("equipment") or md.get("equipment_name")).strip():
                equips = [*equips, _txt(md.get("equipment") or md.get("equipment_name")).strip()]
            activities.append(
                DailyActivity(
                    time=_dt(row.get("signed_at") or row.get("created_at")).astimezone(_SH_TZ),
                    component_uri=comp,
                    pile_id=_txt(md.get("pile_id")) or (comp.rsplit("/", 1)[-1] if comp else ""),
                    location=_txt(md.get("location") or md.get("position")),
                    process_step=_txt(md.get("process_step") or md.get("step_name") or row.get("trip_role")),
                    form_code=_txt(md.get("form_code") or md.get("table_name")),
                    trip_id=_txt(row.get("trip_id")) or (trip_uri.rsplit("/", 1)[-1] if trip_uri else ""),
                    primary_executor=_txt(row.get("executor_name") or md.get("executor_name") or ex.get("holder_name") or ex.get("name")),
                    executor_org=_txt(md.get("executor_org") or ex.get("org_uri")),
                    supervisor=_txt(md.get("supervisor_name") or md.get("supervisor")),
                    equipment_used=[*dict.fromkeys([_txt(x).strip() for x in equips if _txt(x).strip()])],
                    gate_result=_gate_text(gate, "en" if _txt(language).lower() == "en" else "zh"),
                    proof_id=_txt(md.get("proof_id") or row.get("doc_id") or md.get("instance_id")),
                    materials_consumed=mat_by_component.get(comp, []),
                    cost_labor=_round(_num(cost.get("labor"))),
                    cost_equipment=_round(_num(cost.get("equipment"))),
                    cost_material=_round(_num(cost.get("material"))),
                    cost_total=_round(_num(cost.get("labor")) + _num(cost.get("equipment")) + _num(cost.get("material"))),
                    remarks=_txt(md.get("remarks") or md.get("remark")),
                )
            )
        activities.sort(key=lambda x: x.time)

        material_summary = [MaterialSummary(name=k[0], code=k[1], unit=k[2], total_qty=_round(v["qty"]), total_cost=_round(v["cost"])) for k, v in sorted(mat_sum.items(), key=lambda x: x[0][0])]
        equipment_summary = [EquipmentSummary(name=name, shifts=_round(sum((item.cost_equipment / 8000.0) for item in activities if name in item.equipment_used)), hours=_round(sum((item.cost_equipment / 1000.0) for item in activities if name in item.equipment_used)), cost=_round(sum(item.cost_equipment for item in activities if name in item.equipment_used))) for name in sorted({eq for item in activities for eq in item.equipment_used})]
        person_keys = sorted({(item.primary_executor, item.process_step) for item in activities if item.primary_executor and item.primary_executor != "SYSTEM"}, key=lambda x: x[0])
        personnel_summary = [PersonnelSummary(name=n, role=r, hours=_round(sum((item.cost_labor / 280.0) for item in activities if item.primary_executor == n and item.process_step == r)), cost=_round(sum(item.cost_labor for item in activities if item.primary_executor == n and item.process_step == r))) for (n, r) in person_keys]

        progress_summary = ProgressSummary(completed_steps=sum(1 for item in activities if item.proof_id and not _is_gate_fail(item.gate_result)), generated_proofs=len({item.proof_id for item in activities if item.proof_id}), components_completed=sum(1 for snap in snapshot.values() if _txt(_dict(snap).get("status")).lower() == "completed"), components_in_progress=sum(1 for snap in snapshot.values() if _txt(_dict(snap).get("status")).lower() != "completed"), pending_steps=sum(max(int(_dict(snap).get("total_steps") or 0) - int(_dict(snap).get("completed_steps") or 0), 0) for snap in snapshot.values()))
        total_inspections = sum(1 for item in activities if item.gate_result not in {"", "待定", "Pending"})
        failed = sum(1 for item in activities if _is_gate_fail(item.gate_result))
        quality_summary = QualitySummary(total_inspections=total_inspections, passed=max(total_inspections - failed, 0), failed=failed, pass_rate=0.0 if total_inspections <= 0 else round(((total_inspections - failed) / total_inspections) * 100.0, 2))

        material_cost = max(_round(daily_material), _round(sum(_num(v["cost"]) for v in mat_sum.values())))
        cost_summary = CostSummary(daily_labor=_round(daily_labor), daily_equipment=_round(daily_equipment), daily_material=_round(material_cost), daily_total=_round(daily_labor + daily_equipment + material_cost), cumulative_total=_round(cumulative_total))

        anomalies: list[Anomaly] = []
        for item in activities:
            if _is_gate_fail(item.gate_result):
                anomalies.append(Anomaly(type="gate_failed", severity="high", component_uri=item.component_uri, description=f"{item.process_step}检验不合格", action_required="整改后重新检验"))
        for row in materials:
            sd = _dict(row.get("state_data"))
            actual = _num(sd.get("quantity") or sd.get("actual_qty") or sd.get("used_qty"))
            standard = max(_num(sd.get("standard_qty"), actual), 0.0001)
            if actual > standard * 1.1:
                anomalies.append(Anomaly(type="over_consumption", severity="medium", component_uri=_txt(sd.get("component_uri")), description=f"{_txt(sd.get('material_name') or sd.get('material_code'))}超耗{_round(actual - standard)}{_txt(sd.get('unit'))}", action_required="检查超耗原因"))
        today = datetime.now(_SH_TZ).date()
        for ex in exec_rows:
            name = _txt(ex.get("holder_name") or ex.get("name"))
            for cert in _list(ex.get("certificates")):
                c = _dict(cert)
                vu = _txt(c.get("valid_until")).strip()
                if not vu:
                    continue
                try:
                    days = (_parse_date(vu) - today).days
                except Exception:
                    continue
                if days <= 30:
                    anomalies.append(Anomaly(type="cert_expiring", severity="low", description=f"{name}{_txt(c.get('cert_type'))}将在{days}天后过期", action_required="及时办理续期"))

        log = DailyLog(
            log_date=d,
            project_uri=uri,
            project_name=_txt(p.get("name")),
            contract_section=_txt(p.get("contract_no") or p.get("contract_section")),
            weather=_txt(weather).strip(),
            temperature_range=_txt(temperature_range).strip(),
            wind_level=_txt(wind_level).strip(),
            activities=activities,
            material_summary=material_summary,
            equipment_summary=equipment_summary,
            personnel_summary=personnel_summary,
            progress_summary=progress_summary,
            quality_summary=quality_summary,
            cost_summary=cost_summary,
            anomalies=anomalies,
            process_snapshot=snapshot,
            v_uri=f"{uri}/log/{d}",
            language="en" if _txt(language).lower() == "en" else "zh",
            locked=False,
        )
        log.data_hash = self._log_hash(log)
        self._write_draft_once(project_uri=uri, log=log)
        return log

    async def sign_daily_log(self, *, project_uri: str, log_date: str, executor_uri: str = "", signed_by: str = "", weather: str = "", temperature_range: str = "", wind_level: str = "", language: str = "zh") -> DailyLog:
        log = await self.generate_daily_log(project_uri=project_uri, log_date=log_date, weather=weather, temperature_range=temperature_range, wind_level=wind_level, language=language)
        if log.locked:
            return log
        signer_uri = _txt(executor_uri).strip() or f"{_txt(project_uri).strip().rstrip('/')}/role/project-manager/"
        signer_name = _txt(signed_by).strip()
        if signer_uri and not signer_name:
            rows = self.sb.table("san_executors").select("*").eq("executor_uri", signer_uri).limit(1).execute().data or []
            if rows and isinstance(rows[0], dict):
                signer_name = _txt(rows[0].get("holder_name") or rows[0].get("name"))
        log.signed_by = signer_name or "项目负责人"
        log.signed_at = datetime.now(UTC)
        log.locked = True
        log.data_hash = self._log_hash(log)
        proof_id = f"GP-LOGPEG-{hashlib.sha256(f'{log.project_uri}:{log.log_date}:{log.data_hash}'.encode('utf-8')).hexdigest()[:16].upper()}"
        row = ProofUTXOEngine(self.sb).create(
            proof_id=proof_id,
            owner_uri=log.project_uri,
            project_uri=log.project_uri,
            proof_type="document",
            result="PASS",
            state_data={"proof_kind": "logpeg_daily_signed", "log_date": log.log_date, "date": log.log_date, "log": log.model_dump(mode="json"), "executor_uri": signer_uri},
            norm_uri="v://norm/NormPeg/LogPeg/1.0",
            segment_uri=f"{log.project_uri}/log/{log.log_date}",
            signer_uri=signer_uri,
            signer_role="OWNER",
        )
        log.sign_proof = _txt(_dict(row).get("proof_id") or proof_id)
        return log

    def _aggregate(self, logs: list[DailyLog]) -> AggregateSummary:
        total_inspections = sum(item.quality_summary.total_inspections for item in logs)
        total_passed = sum(item.quality_summary.passed for item in logs)
        return AggregateSummary(
            total_completed_steps=sum(item.progress_summary.completed_steps for item in logs),
            total_generated_proofs=sum(item.progress_summary.generated_proofs for item in logs),
            total_pending_steps=sum(item.progress_summary.pending_steps for item in logs),
            total_failed=sum(item.quality_summary.failed for item in logs),
            total_material_cost=_round(sum(item.cost_summary.daily_material for item in logs)),
            total_labor_cost=_round(sum(item.cost_summary.daily_labor for item in logs)),
            total_equipment_cost=_round(sum(item.cost_summary.daily_equipment for item in logs)),
            total_cost=_round(sum(item.cost_summary.daily_total for item in logs)),
            total_components_completed=sum(item.progress_summary.components_completed for item in logs),
            total_components_in_progress=sum(item.progress_summary.components_in_progress for item in logs),
            average_pass_rate=0.0 if total_inspections <= 0 else round((total_passed / total_inspections) * 100.0, 2),
        )

    async def generate_weekly_log(self, project_uri: str, week_start: str, language: str = "zh") -> WeeklyLog:
        start = _parse_date(week_start)
        logs = [await self.generate_daily_log(project_uri, (start + timedelta(days=i)).isoformat(), language=language) for i in range(7)]
        return WeeklyLog(project_uri=_txt(project_uri).rstrip("/"), week_start=start.isoformat(), week_end=(start + timedelta(days=6)).isoformat(), daily_logs=logs, weekly_summary=self._aggregate(logs), language="en" if _txt(language).lower() == "en" else "zh")

    async def generate_monthly_log(self, project_uri: str, year_month: str, language: str = "zh") -> MonthlyLog:
        y, m = datetime.strptime(_txt(year_month).strip(), "%Y-%m").year, datetime.strptime(_txt(year_month).strip(), "%Y-%m").month
        count = calendar.monthrange(y, m)[1]
        logs = [await self.generate_daily_log(project_uri, date(y, m, d).isoformat(), language=language) for d in range(1, count + 1)]
        return MonthlyLog(project_uri=_txt(project_uri).rstrip("/"), month=f"{y:04d}-{m:02d}", daily_logs=logs, monthly_summary=self._aggregate(logs), language="en" if _txt(language).lower() == "en" else "zh")

    async def export_daily_log(self, *, project_uri: str, log_date: str, format: str = "pdf", language: str = "zh") -> tuple[bytes, str, str]:
        log = await self.generate_daily_log(project_uri=project_uri, log_date=log_date, language=language)
        fmt = _txt(format).strip().lower()
        if fmt == "json":
            return json.dumps(log.model_dump(mode="json"), ensure_ascii=False, indent=2).encode("utf-8"), f"logpeg-{_safe_file_token(log.project_name or 'project')}-{log.log_date}.json", "application/json"
        if fmt == "word":
            from docx import Document
            doc = Document()
            doc.add_heading(f"{'施工日志' if log.language == 'zh' else 'Construction Daily Log'} {log.log_date}", level=1)
            doc.add_paragraph(f"Project: {log.project_name} ({log.project_uri})")
            doc.add_paragraph(f"v:// {log.v_uri}")
            table = doc.add_table(rows=1, cols=5)
            for i, h in enumerate(["Time", "Pile", "Step", "Executor", "Result"]):
                table.rows[0].cells[i].text = h
            for item in log.activities:
                row = table.add_row().cells
                row[0].text = item.time.strftime("%H:%M")
                row[1].text = item.pile_id
                row[2].text = item.process_step
                row[3].text = item.primary_executor
                row[4].text = item.gate_result
            buf = BytesIO(); doc.save(buf)
            return buf.getvalue(), f"logpeg-{_safe_file_token(log.project_name or 'project')}-{log.log_date}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        buf = BytesIO(); c = canvas.Canvas(buf, pagesize=A4); w, h = A4; y = h - 40
        c.setFont("Helvetica-Bold", 14); c.drawString(36, y, f"{'施工日志' if log.language == 'zh' else 'Construction Daily Log'} {log.log_date}"); y -= 20
        c.setFont("Helvetica", 10); c.drawString(36, y, f"Project: {log.project_name} ({log.project_uri})"); y -= 14
        c.drawString(36, y, f"Weather:{log.weather or '-'} Temp:{log.temperature_range or '-'} Wind:{log.wind_level or '-'}"); y -= 16
        c.setFont("Helvetica", 9)
        for item in log.activities[:24]: c.drawString(36, y, f"{item.time.strftime('%H:%M')} {item.pile_id} {item.process_step} ({item.gate_result})"[:130]); y -= 12
        c.drawString(36, max(y, 70), f"v:// {log.v_uri}"); c.drawString(36, max(y - 12, 58), f"hash: {log.data_hash}")
        q = qr.QrCodeWidget(log.v_uri); b = q.getBounds(); d = Drawing(80, 80, transform=[80.0 / (b[2] - b[0]), 0, 0, 80.0 / (b[3] - b[1]), 0, 0]); d.add(q); renderPDF.draw(d, c, w - 120, 36)
        c.showPage(); c.save()
        return buf.getvalue(), f"logpeg-{_safe_file_token(log.project_name or 'project')}-{log.log_date}.pdf", "application/pdf"


async def auto_generate_daily_logs(*, sb: Any, date_text: str | None = None) -> LogPegAutoGenerateResult:
    target = _parse_date(date_text or datetime.now(_SH_TZ).date().isoformat()).isoformat()
    rows = sb.table("projects").select("id,v_uri,name,status,enterprise_id").eq("status", "active").order("created_at", desc=False).limit(500).execute().data or []
    engine = LogPegEngine(sb=sb); details: list[dict[str, Any]] = []; generated = failed = 0
    for row in rows:
        if not isinstance(row, dict):
            continue
        uri = _txt(row.get("v_uri")).strip().rstrip("/")
        if not uri:
            continue
        try:
            log = await engine.generate_daily_log(project_uri=uri, log_date=target)
            notify = await _notify_project_manager(sb=sb, project_row=row, log=log, event="logpeg.daily.draft.generated")
            details.append({"project_uri": uri, "project_name": _txt(row.get("name")), "date": target, "status": "generated", "notify": notify}); generated += 1
        except Exception as exc:
            details.append({"project_uri": uri, "project_name": _txt(row.get("name")), "date": target, "status": "failed", "error": f"{exc.__class__.__name__}: {exc}"}); failed += 1
    return LogPegAutoGenerateResult(date=target, generated=generated, failed=failed, details=details)


async def remind_unsigned_daily_logs(*, sb: Any, date_text: str | None = None) -> LogPegAutoGenerateResult:
    target = _parse_date(date_text or (datetime.now(_SH_TZ).date() - timedelta(days=1)).isoformat()).isoformat()
    rows = sb.table("projects").select("id,v_uri,name,status,enterprise_id").eq("status", "active").order("created_at", desc=False).limit(500).execute().data or []
    engine = LogPegEngine(sb=sb); details: list[dict[str, Any]] = []; generated = failed = 0
    for row in rows:
        if not isinstance(row, dict):
            continue
        uri = _txt(row.get("v_uri")).strip().rstrip("/")
        if not uri:
            continue
        try:
            if engine._signed_row(uri, target):
                details.append({"project_uri": uri, "project_name": _txt(row.get("name")), "date": target, "status": "already_signed"}); continue
            log = await engine.generate_daily_log(project_uri=uri, log_date=target)
            notify = await _notify_project_manager(sb=sb, project_row=row, log=log, event="logpeg.daily.signature.reminder")
            details.append({"project_uri": uri, "project_name": _txt(row.get("name")), "date": target, "status": "reminded", "notify": notify}); generated += 1
        except Exception as exc:
            details.append({"project_uri": uri, "project_name": _txt(row.get("name")), "date": target, "status": "failed", "error": f"{exc.__class__.__name__}: {exc}"}); failed += 1
    return LogPegAutoGenerateResult(date=target, generated=generated, failed=failed, details=details)


async def _notify_project_manager(*, sb: Any, project_row: dict[str, Any], log: DailyLog, event: str) -> dict[str, Any]:
    enterprise_id = _txt(project_row.get("enterprise_id")).strip()
    if not enterprise_id:
        return {"attempted": False, "reason": "enterprise_id_missing"}
    cfg_rows = sb.table("enterprise_configs").select("custom_fields").eq("enterprise_id", enterprise_id).limit(1).execute().data or []
    cfg = cfg_rows[0] if cfg_rows and isinstance(cfg_rows[0], dict) else {}
    webhook_url = _txt(_dict(cfg.get("custom_fields")).get("webhook_url")).strip()
    if not webhook_url:
        return {"attempted": False, "reason": "webhook_url_missing"}
    payload = {"event": event, "project_uri": log.project_uri, "project_name": log.project_name, "date": log.log_date, "v_uri": log.v_uri, "data_hash": log.data_hash, "generated_at": datetime.now(UTC).isoformat()}
    try:
        async with httpx.AsyncClient(timeout=8.0) as client:
            response = await client.post(webhook_url, json=payload)
        return {"attempted": True, "success": 200 <= response.status_code < 300, "status_code": response.status_code}
    except Exception as exc:
        return {"attempted": True, "success": False, "reason": f"{exc.__class__.__name__}: {exc}"}


__all__ = ["LogPegEngine", "auto_generate_daily_logs", "remind_unsigned_daily_logs"]
