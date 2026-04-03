"""Component-level DocPeg rendering helpers."""

from __future__ import annotations

import base64
from datetime import datetime, timezone
import io
from pathlib import Path
import re
from typing import Any

import qrcode
from docx import Document
from docx.shared import Inches, Mm
from docxtpl import DocxTemplate, InlineImage

from services.api.domain.execution.triprole_common import (
    as_dict as _as_dict,
    as_list as _as_list,
    to_text as _to_text,
)


def _safe_name(value: str, fallback: str = "component_report") -> str:
    text = _to_text(value).strip()
    text = re.sub(r"[^\w.\-]+", "_", text, flags=re.ASCII).strip("_")
    return text or fallback


def _qr_png_bytes(value: str) -> bytes:
    img = qrcode.make(value)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _build_verify_uri(*, proof_hash: str, verify_base_url: str) -> str:
    normalized_hash = _to_text(proof_hash).strip()
    if not normalized_hash:
        return ""
    base = _to_text(verify_base_url).strip().rstrip("/") or "https://verify.qcspec.com"
    return f"{base}/component-proof/{normalized_hash}"


def build_component_docpeg_context(
    *,
    verification: dict[str, Any],
    verify_base_url: str = "https://verify.qcspec.com",
) -> dict[str, Any]:
    project_uri = _to_text(verification.get("project_uri") or "").strip()
    component_id = _to_text(verification.get("component_id") or "").strip()
    component_uri = _to_text(verification.get("component_uri") or "").strip()
    proof_hash = _to_text(verification.get("proof_hash") or "").strip()
    within_tolerance = bool(verification.get("within_tolerance"))
    verify_uri = _build_verify_uri(proof_hash=proof_hash, verify_base_url=verify_base_url)

    materials = []
    for item in _as_list(verification.get("materials")):
        row = _as_dict(item)
        planned = float(row.get("planned") or 0.0)
        actual = float(row.get("actual") or 0.0)
        deviation_ratio = float(row.get("deviation_ratio") or 0.0)
        materials.append(
            {
                "material_role": _to_text(row.get("material_role") or "").strip(),
                "boq_item_id": _to_text(row.get("boq_item_id") or "").strip(),
                "planned": planned,
                "actual": actual,
                "deviation_ratio": deviation_ratio,
                "yield_strength": _to_text(row.get("yield_strength") or "N/A").strip() or "N/A",
                "norm_result": "PASS" if bool(row.get("within_tolerance")) else "FAIL",
            }
        )

    field_sources = {
        "project_name": "project_uri",
        "component_id": "component_id",
        "material_role": "materials[].material_role",
        "boq_item_id": "materials[].boq_item_id",
        "planned_qty": "materials[].planned",
        "actual_qty": "materials[].actual",
        "deviation_ratio": "materials[].deviation_ratio",
        "yield_strength": "materials[].yield_strength",
        "norm_result": "materials[].norm_result + within_tolerance",
        "proof_hash": "proof_hash",
    }

    return {
        "project_uri": project_uri,
        "component_id": component_id,
        "component_uri": component_uri,
        "kind": _to_text(verification.get("kind") or "").strip(),
        "status": _to_text(verification.get("status") or "").strip(),
        "version": int(verification.get("version") or 1),
        "proof_hash": proof_hash,
        "verify_uri": verify_uri,
        "generated_at": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"),
        "within_tolerance": within_tolerance,
        "total_delta": float(verification.get("total_delta") or 0.0),
        "materials": materials,
        "field_sources": field_sources,
    }


def _render_with_template(
    *,
    template_path: Path,
    context: dict[str, Any],
) -> bytes:
    tpl = DocxTemplate(str(template_path))
    render_ctx = dict(context)
    verify_uri = _to_text(context.get("verify_uri") or "").strip()
    if verify_uri:
        qr_buf = io.BytesIO(_qr_png_bytes(verify_uri))
        render_ctx["proof_qr"] = InlineImage(tpl, qr_buf, width=Mm(25))
    tpl.render(render_ctx, autoescape=False)
    out = io.BytesIO()
    tpl.save(out)
    return out.getvalue()


def _render_fallback_docx(*, context: dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading("Component Conservation Inspection Report", level=1)
    doc.add_paragraph(f"Project: {_to_text(context.get('project_uri') or '-').strip() or '-'}")
    doc.add_paragraph(f"Component: {_to_text(context.get('component_id') or '-').strip() or '-'}")
    doc.add_paragraph(f"Component URI: {_to_text(context.get('component_uri') or '-').strip() or '-'}")
    doc.add_paragraph(f"Kind: {_to_text(context.get('kind') or '-').strip() or '-'}")
    doc.add_paragraph(f"Status: {_to_text(context.get('status') or '-').strip() or '-'}")
    doc.add_paragraph(f"Version: {_to_text(context.get('version') or '-').strip() or '-'}")
    doc.add_paragraph(f"Generated At: {_to_text(context.get('generated_at') or '-').strip() or '-'}")
    doc.add_paragraph(f"Proof Hash: {_to_text(context.get('proof_hash') or '-').strip() or '-'}")
    verify_uri = _to_text(context.get("verify_uri") or "").strip()
    if verify_uri:
        doc.add_paragraph(f"Verify URI: {verify_uri}")
        image = io.BytesIO(_qr_png_bytes(verify_uri))
        doc.add_picture(image, width=Inches(1.4))

    table = doc.add_table(rows=1, cols=7)
    header = table.rows[0].cells
    header[0].text = "Material Role"
    header[1].text = "BOQ Item"
    header[2].text = "Planned Qty"
    header[3].text = "Actual Qty"
    header[4].text = "Deviation"
    header[5].text = "Yield Strength"
    header[6].text = "Norm Result"

    for item in _as_list(context.get("materials")):
        row = _as_dict(item)
        cells = table.add_row().cells
        cells[0].text = _to_text(row.get("material_role") or "-")
        cells[1].text = _to_text(row.get("boq_item_id") or "-")
        cells[2].text = _to_text(row.get("planned") or 0)
        cells[3].text = _to_text(row.get("actual") or 0)
        cells[4].text = f"{float(row.get('deviation_ratio') or 0.0) * 100:.2f}%"
        cells[5].text = _to_text(row.get("yield_strength") or "N/A")
        cells[6].text = _to_text(row.get("norm_result") or "PENDING")

    doc.add_paragraph("Conclusion: " + ("PASS" if bool(context.get("within_tolerance")) else "FAIL"))
    doc.add_paragraph("Total Delta: " + _to_text(context.get("total_delta") or 0.0))
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def render_component_docpeg_docx(
    *,
    context: dict[str, Any],
    template_path: str | Path | None = None,
) -> tuple[bytes, str]:
    resolved_template = Path(template_path).expanduser().resolve() if template_path else None
    if resolved_template and resolved_template.exists():
        return _render_with_template(template_path=resolved_template, context=context), str(resolved_template)
    return _render_fallback_docx(context=context), "fallback_programmatic"


def build_component_docpeg_bundle(
    *,
    verification: dict[str, Any],
    verify_base_url: str = "https://verify.qcspec.com",
    template_path: str | Path | None = None,
    include_docx_base64: bool = True,
) -> dict[str, Any]:
    context = build_component_docpeg_context(
        verification=verification,
        verify_base_url=verify_base_url,
    )
    docx_bytes, template_used = render_component_docpeg_docx(
        context=context,
        template_path=template_path,
    )
    file_name = f"{_safe_name(_to_text(context.get('component_id') or ''), 'component')}_component_report.docx"
    payload: dict[str, Any] = {
        "ok": True,
        "document_type": "component_report",
        "file_name": file_name,
        "template_used": template_used,
        "proof_embedded": bool(_to_text(context.get("proof_hash") or "").strip()),
        "qr_embedded": bool(_to_text(context.get("verify_uri") or "").strip()),
        "docx_size": len(docx_bytes),
        "context": context,
        "field_sources": _as_dict(context.get("field_sources")),
    }
    if include_docx_base64:
        payload["docx_base64"] = base64.b64encode(docx_bytes).decode("ascii")
    return payload


__all__ = [
    "build_component_docpeg_context",
    "render_component_docpeg_docx",
    "build_component_docpeg_bundle",
]
