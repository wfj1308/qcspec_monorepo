"""
DocPeg helpers for BOQ-linked proof chain aggregation and report packaging.
"""

from __future__ import annotations

import base64
from datetime import datetime, timedelta, timezone
import hashlib
import io
import json
import mimetypes
import os
from pathlib import Path
import re
from typing import Any
from urllib import request as urlrequest
import zipfile

import qrcode
from docx import Document
from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage
from services.api.infrastructure.document.pdf_utils import pdf_report_bytes
from services.api.domain.boq.runtime.audit_common import (
    as_dict as _as_dict,
    as_list as _as_list,
    chain_root_hash as _chain_root_hash,
    to_float as _common_to_float,
    to_text as _to_text,
)
from services.api.domain.utxo.integrations import ProofUTXOEngine
from services.api.domain.reporting.runtime.reports_generation import _try_convert_docx_to_pdf
from services.api.domain.verify.runtime.service import get_proof_ancestry, get_proof_descendants


def _to_float(value: Any) -> float | None:
    return _common_to_float(value, regex_fallback=True)


def _parse_dt(value: Any) -> datetime:
    text = _to_text(value).strip()
    if not text:
        return datetime.min
    normalized = text[:-1] + "+00:00" if text.endswith("Z") else text
    try:
        return datetime.fromisoformat(normalized).replace(tzinfo=None)
    except Exception:
        return datetime.min


def _query_rows_by_boq_item_uri(sb: Any, boq_item_uri: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []

    def _extend(data: Any) -> None:
        if isinstance(data, list):
            for row in data:
                if isinstance(row, dict):
                    rows.append(row)

    # 1) fast path: segment_uri is set to boq_item_uri during initialization
    try:
        res = (
            sb.table("proof_utxo")
            .select("*")
            .eq("segment_uri", boq_item_uri)
            .order("created_at", desc=False)
            .limit(1000)
            .execute()
        )
        _extend(res.data)
    except Exception:
        pass

    # 2) JSON path filters
    for key in ("boq_item_uri", "item_uri", "boq_uri"):
        try:
            res = (
                sb.table("proof_utxo")
                .select("*")
                .filter(f"state_data->>{key}", "eq", boq_item_uri)
                .order("created_at", desc=False)
                .limit(1000)
                .execute()
            )
            _extend(res.data)
        except Exception:
            continue

    # 3) local fallback scan
    if not rows:
        try:
            res = (
                sb.table("proof_utxo")
                .select("*")
                .order("created_at", desc=False)
                .limit(5000)
                .execute()
            )
            for row in res.data or []:
                if not isinstance(row, dict):
                    continue
                sd = row.get("state_data") if isinstance(row.get("state_data"), dict) else {}
                if _to_text(sd.get("boq_item_uri") or sd.get("item_uri") or sd.get("boq_uri")) == boq_item_uri:
                    rows.append(row)
        except Exception:
            pass

    dedup: dict[str, dict[str, Any]] = {}
    for row in rows:
        proof_id = _to_text(row.get("proof_id")).strip()
        if not proof_id:
            continue
        dedup[proof_id] = row
    return list(dedup.values())


def get_proof_chain(boq_item_uri: str, sb: Any, *, max_depth: int = 128, max_nodes: int = 512) -> list[dict[str, Any]]:
    """
    Recursively gather the complete proof chain for a BOQ line item.
    Includes ancestor and descendant nodes.
    """
    normalized_uri = _to_text(boq_item_uri).strip()
    if not normalized_uri:
        return []

    seeds = _query_rows_by_boq_item_uri(sb, normalized_uri)
    if not seeds:
        return []

    engine = ProofUTXOEngine(sb)
    out: dict[str, dict[str, Any]] = {}

    for seed in seeds:
        seed_id = _to_text(seed.get("proof_id")).strip()
        if not seed_id:
            continue
        for row in get_proof_ancestry(engine, seed_id, max_depth=max_depth):
            pid = _to_text((row or {}).get("proof_id")).strip()
            if pid:
                out[pid] = row
        for row in get_proof_descendants(engine, seed_id, max_depth=12, max_nodes=max_nodes):
            pid = _to_text((row or {}).get("proof_id")).strip()
            if pid:
                out[pid] = row

    rows = list(out.values())
    rows.sort(key=lambda row: _parse_dt(row.get("created_at")))
    return rows


def _row_fingerprint(row: dict[str, Any]) -> dict[str, Any]:
    sd = row.get("state_data") if isinstance(row.get("state_data"), dict) else {}
    v_uri_refs = [
        _to_text(row.get("project_uri") or "").strip(),
        _to_text(row.get("segment_uri") or "").strip(),
        _to_text(row.get("norm_uri") or "").strip(),
        _to_text(sd.get("v_uri") or "").strip(),
        _to_text(sd.get("boq_item_uri") or sd.get("item_uri") or "").strip(),
        _to_text(sd.get("norm_uri") or sd.get("spec_uri") or "").strip(),
    ]
    v_uri_refs = [uri for uri in v_uri_refs if uri.startswith("v://")]

    payload = {
        "proof_id": _to_text(row.get("proof_id") or ""),
        "proof_hash": _to_text(row.get("proof_hash") or ""),
        "parent_proof_id": _to_text(row.get("parent_proof_id") or ""),
        "project_uri": _to_text(row.get("project_uri") or ""),
        "segment_uri": _to_text(row.get("segment_uri") or ""),
        "proof_type": _to_text(row.get("proof_type") or ""),
        "result": _to_text(row.get("result") or ""),
        "state_data": row.get("state_data") if isinstance(row.get("state_data"), dict) else {},
        "norm_uri": _to_text(row.get("norm_uri") or ""),
    }
    canonical = json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)
    row_hash = hashlib.sha256(canonical.encode("utf-8")).hexdigest()
    return {
        "proof_id": payload["proof_id"],
        "proof_hash": payload["proof_hash"],
        "row_hash": row_hash,
        "parent_proof_id": payload["parent_proof_id"],
        "proof_type": payload["proof_type"],
        "result": payload["result"],
        "created_at": _to_text(row.get("created_at") or ""),
        "segment_uri": payload["segment_uri"],
        "norm_uri": payload["norm_uri"],
        "v_uri_refs": sorted(set(v_uri_refs)),
        "geo_location": _as_dict(sd.get("geo_location")),
        "server_timestamp_proof": _as_dict(sd.get("server_timestamp_proof")),
        "spatiotemporal_anchor_hash": _to_text(sd.get("spatiotemporal_anchor_hash") or "").strip(),
    }


def build_chain_fingerprints(chain_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [_row_fingerprint(row) for row in chain_rows if isinstance(row, dict)]


def build_rebar_inspection_rows(chain_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for idx, row in enumerate(chain_rows, start=1):
        if not isinstance(row, dict):
            continue
        sd = row.get("state_data") if isinstance(row.get("state_data"), dict) else {}
        values = sd.get("values") if isinstance(sd.get("values"), list) else []
        values_num = [v for v in (_to_float(x) for x in values) if v is not None]
        measured = _to_float(sd.get("value"))
        if measured is None and values_num:
            measured = round(sum(values_num) / len(values_num), 4)

        design = _to_float(sd.get("design"))
        if design is None:
            design = _to_float(sd.get("standard"))

        deviation = _to_float(sd.get("deviation_percent"))
        if deviation is None and measured is not None and design is not None and abs(design) > 1e-9:
            deviation = round(((measured - design) / abs(design)) * 100.0, 4)

        fp = _row_fingerprint(row)
        stage = _to_text(sd.get("lifecycle_stage") or sd.get("status") or "").strip().upper()
        trip_action = _to_text(sd.get("trip_action") or "").strip().lower()
        label_seed = _to_text(sd.get("item_name") or sd.get("type_name") or sd.get("type") or row.get("proof_type") or "-")
        if stage in {"ENTRY", "INSTALLATION", "VARIATION", "SETTLEMENT"}:
            label_seed = f"{stage}:{trip_action or label_seed}"
        rows.append(
            {
                "index": idx,
                "item_name": label_seed,
                "design_value": design,
                "measured_value": measured,
                "deviation_percent": deviation,
                "result": _to_text(row.get("result") or sd.get("result") or "PENDING").upper(),
                "proof_id": fp["proof_id"],
                "proof_hash": fp["proof_hash"],
                "row_hash": fp["row_hash"],
                "created_at": _to_text(row.get("created_at") or ""),
                "boq_item_uri": _to_text(sd.get("boq_item_uri") or row.get("segment_uri") or ""),
                "spec_excerpt": _to_text(sd.get("spec_excerpt") or ""),
            }
        )
    return rows


def build_timeline_rows(chain_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    timeline: list[dict[str, Any]] = []
    for idx, row in enumerate(chain_rows, start=1):
        if not isinstance(row, dict):
            continue
        sd = row.get("state_data") if isinstance(row.get("state_data"), dict) else {}
        stage = _to_text(sd.get("lifecycle_stage") or sd.get("status") or "").strip().upper()
        action = _to_text(sd.get("trip_action") or "").strip().lower()
        label = f"{stage}:{action}" if stage else _to_text(row.get("proof_type") or "proof").lower()
        timeline.append(
            {
                "step": idx,
                "label": label,
                "result": _to_text(row.get("result") or "PENDING").upper(),
                "time": _to_text(row.get("created_at") or ""),
                "proof_id": _to_text(row.get("proof_id") or ""),
                "proof_hash": _to_text(row.get("proof_hash") or ""),
                "parent_proof_id": _to_text(row.get("parent_proof_id") or ""),
            }
        )
    return timeline


def _qr_inline(tpl: DocxTemplate, uri: str, *, size_mm: int = 24) -> InlineImage:
    img = qrcode.make(uri)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return InlineImage(tpl, buf, width=Mm(size_mm))


def _scan_confirm_secret() -> str:
    return _to_text(
        os.getenv("QCSPEC_SCAN_CONFIRM_SECRET")
        or os.getenv("SCAN_CONFIRM_SECRET")
        or "qcspec-scan-confirm-v1"
    ).strip()


def _normalize_role(raw: Any) -> str:
    text = _to_text(raw).strip().lower()
    aliases = {
        "施工": "contractor",
        "施工方": "contractor",
        "承包方": "contractor",
        "contractor": "contractor",
        "监理": "supervisor",
        "监理方": "supervisor",
        "supervisor": "supervisor",
        "业主": "owner",
        "业主方": "owner",
        "甲方": "owner",
        "owner": "owner",
    }
    return aliases.get(text) or aliases.get(_to_text(raw).strip()) or text


def _extract_scan_signer(chain_rows: list[dict[str, Any]]) -> tuple[str, str]:
    if not chain_rows:
        return ("", "")
    latest = chain_rows[-1] if isinstance(chain_rows[-1], dict) else {}
    sd = latest.get("state_data") if isinstance(latest.get("state_data"), dict) else {}
    consensus = sd.get("consensus") if isinstance(sd.get("consensus"), dict) else {}
    signatures = consensus.get("signatures") if isinstance(consensus.get("signatures"), list) else []
    if not signatures and isinstance(sd.get("signatures"), list):
        signatures = sd.get("signatures")
    signer_did = ""
    signer_role = ""
    for sig in signatures:
        if not isinstance(sig, dict):
            continue
        did = _to_text(sig.get("did") or "").strip()
        role = _normalize_role(sig.get("role"))
        if did.startswith("did:") and role == "contractor":
            return did, "contractor"
        if did.startswith("did:") and not signer_did:
            signer_did = did
            signer_role = role or _to_text(sig.get("role") or "").strip().lower()
    return signer_did, signer_role


def _build_scan_confirm_token(*, proof_id: str, signer_did: str, signer_role: str) -> tuple[dict[str, Any], str]:
    issued_at = datetime.now(timezone.utc).isoformat()
    expires_at = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
    payload: dict[str, Any] = {
        "proof_id": _to_text(proof_id).strip(),
        "signer_did": _to_text(signer_did).strip(),
        "signer_role": _normalize_role(signer_role) or _to_text(signer_role).strip().lower(),
        "issued_at": issued_at,
        "expires_at": expires_at,
        "nonce": hashlib.sha256(f"{proof_id}|{issued_at}|{signer_did}".encode("utf-8")).hexdigest()[:24],
    }
    canonical = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    payload["token_hash"] = hashlib.sha256(f"{canonical}|{_scan_confirm_secret()}".encode("utf-8")).hexdigest()
    token = base64.urlsafe_b64encode(
        json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    ).decode("utf-8").rstrip("=")
    return payload, token


def build_rebar_report_context(
    *,
    boq_item_uri: str,
    chain_rows: list[dict[str, Any]],
    project_meta: dict[str, Any],
    verify_base_url: str = "https://verify.qcspec.com",
) -> dict[str, Any]:
    table_rows = build_rebar_inspection_rows(chain_rows)
    fingerprints = build_chain_fingerprints(chain_rows)
    timeline = build_timeline_rows(chain_rows)

    last_proof_id = _to_text(table_rows[-1].get("proof_id") if table_rows else "")
    verify_uri = ""
    if last_proof_id:
        verify_uri = f"{verify_base_url.rstrip('/')}/v/{last_proof_id}?trace=true"
    scan_confirm_payload: dict[str, Any] = {}
    scan_confirm_token = ""
    scan_confirm_uri = ""
    scan_signer_did = ""
    scan_signer_role = ""
    if last_proof_id:
        scan_signer_did, scan_signer_role = _extract_scan_signer(chain_rows)
        if scan_signer_did.startswith("did:"):
            scan_confirm_payload, scan_confirm_token = _build_scan_confirm_token(
                proof_id=last_proof_id,
                signer_did=scan_signer_did,
                signer_role=scan_signer_role or "contractor",
            )
            scan_confirm_uri = f"{verify_base_url.rstrip('/')}/v/{last_proof_id}?scan_token={scan_confirm_token}"

    pass_count = sum(1 for row in table_rows if row.get("result") == "PASS")
    fail_count = sum(1 for row in table_rows if row.get("result") == "FAIL")

    return {
        "project_name": _to_text(project_meta.get("project_name") or project_meta.get("name") or "-"),
        "project_uri": _to_text(project_meta.get("project_uri") or ""),
        "contract_no": _to_text(project_meta.get("contract_no") or "-"),
        "boq_item_uri": boq_item_uri,
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
        "records": table_rows,
        "timeline": timeline,
        "proof_chain": fingerprints,
        "chain_count": len(fingerprints),
        "chain_root_hash": _chain_root_hash(fingerprints),
        "proof_id": last_proof_id,
        "verify_uri": verify_uri,
        "total_count": len(table_rows),
        "pass_count": pass_count,
        "fail_count": fail_count,
        "summary_result_cn": "不合格" if fail_count > 0 else "合格",
        "total_proof_hash": _chain_root_hash(fingerprints),
        "artifact_uri": _to_text(project_meta.get("artifact_uri") or ""),
        "gitpeg_anchor": _to_text(project_meta.get("gitpeg_anchor") or ""),
        "scan_confirm_uri": scan_confirm_uri,
        "scan_confirm_token": scan_confirm_token,
        "scan_confirm_payload": scan_confirm_payload,
        "scan_confirm_signer_did": scan_signer_did,
        "scan_confirm_signer_role": scan_signer_role,
    }


def render_rebar_inspection_docx(
    *,
    template_path: str | Path,
    context: dict[str, Any],
) -> bytes:
    path = Path(template_path).expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(f"Docx template not found: {path}")

    tpl = DocxTemplate(str(path))
    render_ctx = dict(context)
    verify_uri = _to_text(render_ctx.get("verify_uri") or "").strip()
    if verify_uri:
        qr = _qr_inline(tpl, verify_uri)
        render_ctx["qr_image"] = qr
        render_ctx["verify_qr"] = qr
    scan_confirm_uri = _to_text(render_ctx.get("scan_confirm_uri") or "").strip()
    if scan_confirm_uri:
        render_ctx["scan_confirm_qr"] = _qr_inline(tpl, scan_confirm_uri)

    tpl.render(render_ctx, autoescape=False)
    buf = io.BytesIO()
    tpl.save(buf)
    rendered = buf.getvalue()
    hierarchy_rows = render_ctx.get("hierarchy_summary_rows")
    has_hierarchy = isinstance(hierarchy_rows, list) and bool(hierarchy_rows)
    evidence_count = render_ctx.get("evidence_count")
    evidence_source = _to_text(render_ctx.get("evidence_source") or "").strip()
    total_proof_hash = _to_text(render_ctx.get("total_proof_hash") or "").strip()
    risk_score = _to_text(render_ctx.get("risk_score") or "").strip()
    risk_issue_count = _to_text(render_ctx.get("risk_issue_count") or "").strip()
    asset_origin_statement = _to_text(render_ctx.get("asset_origin_statement") or "").strip()
    sealing_trip = _as_dict(render_ctx.get("sealing_trip"))
    sealing_pattern_id = _to_text(sealing_trip.get("pattern_id") or "").strip()
    sealing_symmetry = _to_text(sealing_trip.get("symmetry") or "").strip()
    sealing_scan_hint = _to_text(sealing_trip.get("scan_hint") or "").strip()
    has_evidence_summary = bool(
        evidence_count is not None
        or evidence_source
        or total_proof_hash
        or risk_score
        or risk_issue_count
        or asset_origin_statement
        or sealing_pattern_id
    )
    if not has_hierarchy and not has_evidence_summary:
        return rendered
    try:
        doc = Document(io.BytesIO(rendered))
        if has_evidence_summary:
            doc.add_paragraph("证据链摘要")
            if evidence_count is not None:
                doc.add_paragraph(f"证据数量: {evidence_count}")
            if evidence_source:
                doc.add_paragraph(f"证据来源: {evidence_source}")
            if total_proof_hash:
                doc.add_paragraph(f"Total Proof Hash: {total_proof_hash}")
            if risk_score:
                doc.add_paragraph(f"风险分值: {risk_score}")
            if risk_issue_count:
                doc.add_paragraph(f"风险项数量: {risk_issue_count}")
            if asset_origin_statement:
                doc.add_paragraph(f"量值来源说明: {asset_origin_statement}")
            if sealing_pattern_id:
                doc.add_paragraph("主权防伪底纹")
                doc.add_paragraph(f"Pattern ID: {sealing_pattern_id}")
                if sealing_symmetry:
                    doc.add_paragraph(f"对称策略: {sealing_symmetry}")
                if sealing_scan_hint:
                    doc.add_paragraph(f"验真码: {sealing_scan_hint}")
        verify_uri = _to_text(render_ctx.get("verify_uri") or "").strip()
        scan_confirm_uri = _to_text(render_ctx.get("scan_confirm_uri") or "").strip()
        if verify_uri or scan_confirm_uri:
            doc.add_paragraph("验真摘要")
            if verify_uri:
                doc.add_paragraph(f"验真链接: {verify_uri}")
            if scan_confirm_uri:
                doc.add_paragraph(f"复核链接: {scan_confirm_uri}")
        if not has_hierarchy:
            out = io.BytesIO()
            doc.save(out)
            return out.getvalue()
        doc.add_paragraph("分部分项汇总（章/节/目/细目）")
        table = doc.add_table(rows=1, cols=6)
        head = table.rows[0].cells
        head[0].text = "层级"
        head[1].text = "编码"
        head[2].text = "类型"
        head[3].text = "名称"
        head[4].text = "已结算/设计"
        head[5].text = "进度(%)"
        for row in hierarchy_rows:
            if not isinstance(row, dict):
                continue
            depth = int(_to_float(row.get("depth")) or 0)
            code = _to_text(row.get("code") or "").strip()
            node_type = _to_text(row.get("node_type") or "").strip()
            name = _to_text(row.get("item_name") or "").strip()
            settled = _to_text(row.get("settled_quantity") or 0).strip()
            design = _to_text(row.get("design_quantity") or 0).strip()
            pct = _to_text(row.get("progress_percent") or 0).strip()
            cells = table.add_row().cells
            cells[0].text = str(depth)
            cells[1].text = code
            cells[2].text = node_type
            cells[3].text = name
            cells[4].text = f"{settled} / {design}"
            cells[5].text = pct
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception:
        return rendered


def render_rebar_inspection_pdf(*, docx_bytes: bytes, context: dict[str, Any]) -> bytes:
    converted = _try_convert_docx_to_pdf(docx_bytes)
    if converted:
        return converted

    lines = [
        f"QCSpec Rebar Inspection Report: {_to_text(context.get('project_name') or '-')}",
        f"BOQ Item URI: {_to_text(context.get('boq_item_uri') or '-')}",
        f"Proof Chain Count: {_to_text(context.get('chain_count') or 0)}",
        f"Chain Root Hash: {_to_text(context.get('chain_root_hash') or '-')}",
        f"Total Proof Hash: {_to_text(context.get('total_proof_hash') or context.get('chain_root_hash') or '-')}",
        f"Evidence Count: {_to_text(context.get('evidence_count') or '-')}",
        f"Evidence Source: {_to_text(context.get('evidence_source') or '-')}",
        f"Artifact URI: {_to_text(context.get('artifact_uri') or '-')}",
        f"GitPeg Anchor: {_to_text(context.get('gitpeg_anchor') or '-')}",
        f"Verify URI: {_to_text(context.get('verify_uri') or '-')}",
        f"Scan Confirm URI: {_to_text(context.get('scan_confirm_uri') or '-')}",
        f"Scan Signer DID: {_to_text(context.get('scan_confirm_signer_did') or '-')}",
        f"Credit Score: {_to_text(context.get('credit_score') or '-')}",
        f"Credit Grade: {_to_text(context.get('credit_grade') or '-')}",
        f"Credit Fast Track: {_to_text(context.get('credit_fast_track_eligible') or '-')}",
        f"Geo Trust Level: {_to_text(context.get('trust_level') or '-')}",
        f"Geo Fence Warning: {_to_text(context.get('geo_fence_warning') or '-')}",
        f"Sensor Device SN: {_to_text(context.get('sensor_device_sn') or '-')}",
        f"Sensor Calibration Valid Until: {_to_text(context.get('sensor_calibration_valid_until') or '-')}",
        f"Sensor Calibration Valid: {_to_text(context.get('sensor_calibration_valid') or '-')}",
        f"Biometric OK: {_to_text(context.get('biometric_ok') or '-')}",
        f"Biometric Verified/Required: {_to_text(context.get('biometric_verified_count') or '-')}/{_to_text(context.get('biometric_required_count') or '-')}",
        f"Risk Score: {_to_text(context.get('risk_score') or '-')}",
        f"Risk Issue Count: {_to_text(context.get('risk_issue_count') or '-')}",
        f"DID Reputation Score: {_to_text(_as_dict(_as_dict(context.get('risk_audit')).get('did_reputation')).get('aggregate_score') or '-')}",
        f"DID Sampling Multiplier: {_to_text(_as_dict(_as_dict(context.get('risk_audit')).get('did_reputation')).get('sampling_multiplier') or '-')}",
        f"Asset Origin: {_to_text(context.get('asset_origin_statement') or '-')}",
        f"Sealing Pattern ID: {_to_text(_as_dict(context.get('sealing_trip')).get('pattern_id') or '-')}",
        f"Sealing Symmetry: {_to_text(_as_dict(context.get('sealing_trip')).get('symmetry') or '-')}",
        f"Sealing Scan Hint: {_to_text(_as_dict(context.get('sealing_trip')).get('scan_hint') or '-')}",
        f"Hierarchy Root Hash: {_to_text(context.get('hierarchy_root_hash') or '-')}",
        f"Chapter Progress(%): {_to_text(_as_dict(context.get('chapter_progress')).get('progress_percent') or '-')}",
        "---- Rows ----",
    ]
    for row in context.get("records") if isinstance(context.get("records"), list) else []:
        if not isinstance(row, dict):
            continue
        lines.append(
            f"{_to_text(row.get('index'))}. {_to_text(row.get('item_name'))} | design={_to_text(row.get('design_value'))} | measured={_to_text(row.get('measured_value'))} | result={_to_text(row.get('result'))} | row_hash={_to_text(row.get('row_hash'))}"
        )
    lines.append("---- Timeline ----")
    for item in context.get("timeline") if isinstance(context.get("timeline"), list) else []:
        if not isinstance(item, dict):
            continue
        lines.append(
            f"{_to_text(item.get('step'))}. {_to_text(item.get('label'))} | result={_to_text(item.get('result'))} | time={_to_text(item.get('time'))} | proof={_to_text(item.get('proof_id'))}"
        )
    hierarchy_rows = context.get("hierarchy_summary_rows")
    if isinstance(hierarchy_rows, list) and hierarchy_rows:
        lines.append("---- Hierarchy Summary ----")
        for row in hierarchy_rows:
            if not isinstance(row, dict):
                continue
            indent = "  " * max(0, int(_to_float(row.get("depth")) or 1) - 1)
            lines.append(
                f"{indent}{_to_text(row.get('code') or '-')} [{_to_text(row.get('node_type') or '-')}] | progress={_to_text(row.get('progress_percent') or 0)}% | settled={_to_text(row.get('settled_quantity') or 0)} | design={_to_text(row.get('design_quantity') or 0)}"
            )
    seal_ascii = _as_list(_as_dict(context.get("sealing_trip")).get("ascii_pattern"))
    if seal_ascii:
        lines.append("---- Sovereign Sealing Pattern ----")
        for row in seal_ascii[:14]:
            lines.append(_to_text(row))

    return pdf_report_bytes(
        lines,
        watermark=_to_text(_as_dict(context.get("sealing_trip")).get("watermark") or "QCSpec BOQ Proof Chain"),
        trust_ok=True,
    )


def _safe_file_name(name: str, fallback: str) -> str:
    raw = _to_text(name).strip() or fallback
    safe = re.sub(r"[^\w.\-]+", "_", raw, flags=re.ASCII)
    return safe[:180] or fallback


def _fetch_binary(url: str, timeout: float = 10.0) -> bytes | None:
    target = _to_text(url).strip()
    if not target or not (target.startswith("http://") or target.startswith("https://")):
        return None
    req = urlrequest.Request(target, method="GET", headers={"User-Agent": "QCSpec-DocPeg/1.0"})
    try:
        with urlrequest.urlopen(req, timeout=timeout) as resp:
            return resp.read()
    except Exception:
        return None


def _extract_evidence_from_chain(proof_chain: list[dict[str, Any]]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for row in proof_chain:
        if not isinstance(row, dict):
            continue
        sd = row.get("state_data") if isinstance(row.get("state_data"), dict) else {}
        evidence = sd.get("evidence") if isinstance(sd.get("evidence"), list) else []
        for item in evidence:
            if not isinstance(item, dict):
                continue
            key = _to_text(item.get("id") or item.get("url") or item.get("file_name") or "").strip()
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(
                {
                    "id": _to_text(item.get("id") or ""),
                    "file_name": _to_text(item.get("file_name") or ""),
                    "url": _to_text(item.get("url") or ""),
                    "content_type": _to_text(item.get("media_type") or item.get("content_type") or ""),
                    "proof_id": _to_text(item.get("proof_id") or row.get("proof_id") or ""),
                }
            )
    return out


def build_dsp_zip_package(
    *,
    report_pdf_bytes: bytes,
    docx_bytes: bytes | None,
    proof_chain: list[dict[str, Any]],
    context: dict[str, Any],
    evidence_items: list[dict[str, Any]] | None = None,
) -> bytes:
    fingerprints = build_chain_fingerprints(proof_chain)
    provenance_chain = {
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
        "boq_item_uri": _to_text(context.get("boq_item_uri") or ""),
        "project_uri": _to_text(context.get("project_uri") or ""),
        "verify_uri": _to_text(context.get("verify_uri") or ""),
        "nodes": fingerprints,
        "v_uri_refs": sorted(
            {
                uri
                for row in fingerprints
                for uri in (row.get("v_uri_refs") if isinstance(row.get("v_uri_refs"), list) else [])
                if _to_text(uri).startswith("v://")
            }
        ),
    }
    fingerprint_json = {
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
        "proof_chain": fingerprints,
        "chain_root_hash": _chain_root_hash(fingerprints),
        "context": {
            "project_name": _to_text(context.get("project_name") or ""),
            "project_uri": _to_text(context.get("project_uri") or ""),
            "boq_item_uri": _to_text(context.get("boq_item_uri") or ""),
            "proof_id": _to_text(context.get("proof_id") or ""),
            "verify_uri": _to_text(context.get("verify_uri") or ""),
        },
    }
    spatiotemporal_anchors = [
        {
            "proof_id": _to_text(row.get("proof_id") or "").strip(),
            "spatiotemporal_anchor_hash": _to_text(row.get("spatiotemporal_anchor_hash") or "").strip(),
            "geo_location": _as_dict(row.get("geo_location")),
            "server_timestamp_proof": _as_dict(row.get("server_timestamp_proof")),
            "created_at": _to_text(row.get("created_at") or "").strip(),
            "trip_action": _to_text(row.get("trip_action") or "").strip().lower(),
        }
        for row in fingerprints
        if _to_text(row.get("spatiotemporal_anchor_hash") or "").strip()
    ]
    spatiotemporal_anchor_json = {
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
        "boq_item_uri": _to_text(context.get("boq_item_uri") or ""),
        "project_uri": _to_text(context.get("project_uri") or ""),
        "anchor_count": len(spatiotemporal_anchors),
        "anchors": spatiotemporal_anchors,
    }
    spatiotemporal_anchor_json["anchors_hash"] = hashlib.sha256(
        json.dumps(spatiotemporal_anchor_json, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")
    ).hexdigest()
    fingerprint_json["spatiotemporal_anchor_hash"] = spatiotemporal_anchor_json["anchors_hash"]
    fingerprint_json["spatiotemporal_anchor_count"] = len(spatiotemporal_anchors)
    signature_payload = json.dumps(
        {
            "chain_root_hash": fingerprint_json["chain_root_hash"],
            "proof_id": _to_text(context.get("proof_id") or ""),
            "project_uri": _to_text(context.get("project_uri") or ""),
            "boq_item_uri": _to_text(context.get("boq_item_uri") or ""),
        },
        ensure_ascii=False,
        sort_keys=True,
        default=str,
    )
    signature = hashlib.sha256(signature_payload.encode("utf-8")).hexdigest()
    signature_json = {
        "algorithm": "SHA256-MOCK-SIGNATURE",
        "signed_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
        "signature": signature,
        "signer_uri": _to_text(context.get("project_uri") or "v://executor/system"),
        "payload_hash": hashlib.sha256(signature_payload.encode("utf-8")).hexdigest(),
    }
    sealing_trip = _as_dict(context.get("sealing_trip"))

    if evidence_items is None:
        evidence_items = _extract_evidence_from_chain(proof_chain)

    evidence_manifest: list[dict[str, Any]] = []
    evidence_files: list[tuple[str, bytes]] = []

    for idx, item in enumerate(evidence_items or [], start=1):
        if not isinstance(item, dict):
            continue
        file_name = _safe_file_name(_to_text(item.get("file_name") or ""), f"evidence_{idx}.bin")

        blob: bytes | None = None
        if isinstance(item.get("content"), (bytes, bytearray)):
            blob = bytes(item.get("content"))
        elif _to_text(item.get("url")):
            blob = _fetch_binary(_to_text(item.get("url")))

        if blob is None:
            evidence_manifest.append(
                {
                    "index": idx,
                    "file_name": file_name,
                    "status": "unavailable",
                    "source_url": _to_text(item.get("url") or ""),
                }
            )
            continue

        ext = mimetypes.guess_extension(_to_text(item.get("content_type") or "")) or ""
        if "." not in file_name and ext:
            file_name = f"{file_name}{ext}"

        sha = hashlib.sha256(blob).hexdigest()
        evidence_manifest.append(
            {
                "index": idx,
                "file_name": file_name,
                "status": "ok",
                "size": len(blob),
                "sha256": sha,
                "source_url": _to_text(item.get("url") or ""),
            }
        )
        evidence_files.append((f"evidence/{file_name}", blob))

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("report.pdf", report_pdf_bytes)
        if docx_bytes is not None:
            zf.writestr("report.docx", docx_bytes)
        zf.writestr("proof_chain.json", json.dumps(proof_chain, ensure_ascii=False, indent=2, default=str))
        zf.writestr("provenance_chain.json", json.dumps(provenance_chain, ensure_ascii=False, indent=2, default=str))
        zf.writestr("fingerprint.json", json.dumps(fingerprint_json, ensure_ascii=False, indent=2, default=str))
        zf.writestr(
            "spatiotemporal_anchor.json",
            json.dumps(spatiotemporal_anchor_json, ensure_ascii=False, indent=2, default=str),
        )
        zf.writestr("signature.json", json.dumps(signature_json, ensure_ascii=False, indent=2, default=str))
        if sealing_trip:
            zf.writestr("sealing_trip.json", json.dumps(sealing_trip, ensure_ascii=False, indent=2, default=str))
        zf.writestr("evidence/manifest.json", json.dumps(evidence_manifest, ensure_ascii=False, indent=2, default=str))
        for path, blob in evidence_files:
            zf.writestr(path, blob)

    buf.seek(0)
    return buf.read()

