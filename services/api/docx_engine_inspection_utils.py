"""
Inspection-focused helpers extracted from DocxEngine.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document


def resolve_test_type(
    engine: Any,
    state_data: dict[str, Any],
    *,
    fallback_type: str,
    road_flatness_cn: str,
    rebar_spacing_cn: str,
) -> tuple[str, str]:
    raw_type = engine._to_text(state_data.get("test_type") or state_data.get("type") or "").strip()
    raw_name = engine._to_text(
        state_data.get("test_name")
        or state_data.get("type_name")
        or state_data.get("inspection_item")
        or ""
    ).strip()

    if not raw_name and raw_type:
        raw_name = raw_type
    if not raw_type:
        raw_type = raw_name or fallback_type

    probe = {"type": raw_type, "type_name": raw_name or raw_type}
    if not raw_name:
        if engine._is_flatness_like(probe):
            raw_name = road_flatness_cn
        elif raw_type == "rebar_spacing":
            raw_name = rebar_spacing_cn
        else:
            raw_name = raw_type
    return raw_type, raw_name


def extract_executor_name(engine: Any, signing: dict[str, Any], *, fallback_uri: str) -> str:
    name = engine._to_text(
        signing.get("executor_name")
        or signing.get("name")
        or signing.get("signer_name")
        or signing.get("display_name")
        or ""
    ).strip()
    if name:
        return name
    uri = engine._to_text(fallback_uri).strip().rstrip("/")
    if not uri:
        return "-"
    tail = uri.split("/")[-1].strip()
    return tail or "-"


def extract_executor_id(engine: Any, signing: dict[str, Any], *, fallback_uri: str, fallback_name: str) -> str:
    for key in ("executor_id", "executor_uid", "uid", "id"):
        val = engine._to_text(signing.get(key)).strip()
        if val:
            return val
    uri = engine._to_text(fallback_uri).strip().rstrip("/")
    if uri:
        return uri.split("/")[-1].strip() or engine._to_text(fallback_name).strip() or "unknown"
    return engine._to_text(fallback_name).strip() or "unknown"


def build_v_uri_tree(
    engine: Any,
    *,
    project_uri: str,
    segment_uri: str,
    v_uri: str,
    proof_id: str,
    stake: str,
    verify_uri: str,
) -> dict[str, Any]:
    return {
        "project_uri": engine._to_text(project_uri),
        "segment_uri": engine._to_text(segment_uri),
        "v_uri": engine._to_text(v_uri),
        "proof_id": engine._to_text(proof_id),
        "stake": engine._to_text(stake),
        "verify_uri": engine._to_text(verify_uri),
        "nodes": [
            {"name": "project", "uri": engine._to_text(project_uri)},
            {"name": "segment", "uri": engine._to_text(segment_uri), "stake": engine._to_text(stake)},
            {"name": "utxo", "uri": engine._to_text(v_uri), "proof_id": engine._to_text(proof_id)},
        ],
    }


def normalize_inspection_proof(
    engine: Any,
    proof: dict[str, Any],
    *,
    standard_op_plus_minus: str,
    schema_mode_value_standard_min: str,
) -> dict[str, Any]:
    if not isinstance(proof, dict):
        return proof
    out = dict(proof)
    sd = out.get("state_data") if isinstance(out.get("state_data"), dict) else {}
    sd_out = dict(sd)
    test_type, test_type_name = engine._resolve_test_type(sd_out, fallback_type="inspection")
    schema_mode = engine._resolve_schema_mode(sd_out, test_type=test_type, test_type_name=test_type_name)
    sd_out.setdefault("test_type", test_type)
    sd_out.setdefault("type", test_type)
    sd_out.setdefault("test_name", test_type_name)
    sd_out.setdefault("type_name", test_type_name)
    sd_out.setdefault("schema_mode", schema_mode)
    if sd_out.get("standard_value") is None:
        if sd_out.get("standard") is not None:
            sd_out["standard_value"] = sd_out.get("standard")
        elif sd_out.get("design") is not None:
            sd_out["standard_value"] = sd_out.get("design")
    if not engine._to_text(sd_out.get("standard_op")).strip():
        probe = {"type": test_type, "type_name": test_type_name}
        if sd_out.get("limit") not in (None, "", "-"):
            sd_out["standard_op"] = standard_op_plus_minus
        elif engine._is_compaction_like(probe) or schema_mode == schema_mode_value_standard_min:
            sd_out["standard_op"] = ">="
        else:
            sd_out["standard_op"] = "<="
    if sd_out.get("value") is None and isinstance(sd_out.get("values"), list) and len(sd_out["values"]) == 1:
        sd_out["value"] = sd_out["values"][0]
    out["state_data"] = sd_out
    return out


def rewrite_inspection_table_rows(
    engine: Any,
    doc_bytes: bytes,
    rows: list[dict[str, Any]],
    *,
    header_labels: list[str] | None = None,
    fail_cn: str,
    fail_rgb: Any,
) -> bytes:
    """
    Rewrite inspection detail table rows so exported table is 1:1 with records.
    Avoid legacy static rows (e.g. fixed rebar placeholders / ??? cell).
    """
    if not rows:
        return doc_bytes
    src = io.BytesIO(doc_bytes)
    doc = Document(src)
    if len(doc.tables) < 2:
        return doc_bytes

    table = doc.tables[1]
    if len(table.rows) < 1:
        return doc_bytes

    labels = (header_labels or ["检查项目", "检查项目", "规范要求", "设计值", "实测值", "判定"])[:6]
    for idx in range(min(len(table.rows[0].cells), 6)):
        table.rows[0].cells[idx].text = labels[idx]

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for row in rows:
        cells = table.add_row().cells
        test_name = engine._to_text(row.get("test_type_name") or row.get("test_type") or "-")
        unit = engine._to_text(row.get("unit") or "").strip()
        test_label = f"{test_name} ({unit})" if unit and unit != "-" else test_name
        sub_item = engine._to_text(row.get("stake") or row.get("location") or "-")
        limit = engine._to_text(row.get("norm_requirement") or row.get("limit") or "-")
        standard = engine._to_text(row.get("standard") or row.get("design") or "-")
        measured = engine._values_single_line(row)
        result_cn = engine._to_text(row.get("result_cn") or "-")
        raw_result_code = engine._to_text(row.get("result") or "").upper()
        proof_hash = engine._to_text(row.get("proof_hash") or "")
        if proof_hash:
            measured = f"{measured}\nPF:{proof_hash[:20]}..."
        deviation_pct = row.get("deviation_percent")
        if deviation_pct is not None:
            try:
                deviation_text = f"{float(deviation_pct):+0.2f}%"
                result_cn = f"{result_cn} ({deviation_text})"
            except Exception:
                pass
        is_fail = raw_result_code == "FAIL" or engine._to_text(row.get("result_cn") or "") == fail_cn

        engine._set_cell_text(cells[0], test_label or "-")
        engine._set_cell_text(cells[1], sub_item or "-")
        engine._set_cell_text(cells[2], limit or "-")
        engine._set_cell_text(cells[3], standard or "-")
        engine._set_cell_text(cells[4], measured or "-", color=fail_rgb if is_fail else None)
        engine._set_cell_text(cells[5], result_cn or "-", color=fail_rgb if is_fail else None)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
