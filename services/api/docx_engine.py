"""
QCSpec DOCX engine for sovereign report rendering.
"""

from __future__ import annotations

import io
import os
from datetime import datetime
from pathlib import Path
from typing import Any

import qrcode
from docx.shared import Mm, RGBColor
from docxtpl import DocxTemplate, InlineImage, Listing
from services.api.docx_engine_inspection_utils import (
    build_v_uri_tree as _build_v_uri_tree_util,
    extract_executor_id as _extract_executor_id_util,
    extract_executor_name as _extract_executor_name_util,
    normalize_inspection_proof as _normalize_inspection_proof_util,
    resolve_test_type as _resolve_test_type_util,
    rewrite_inspection_table_rows as _rewrite_inspection_table_rows_util,
)
from services.api.docx_engine_mock_data import build_rebar_live_mock_case
from services.api.docx_engine_report_utils import (
    build_docpeg_context as _build_docpeg_context_util,
    build_docpeg_row as _build_docpeg_row_util,
)
from services.api.docx_engine_rule_utils import (
    build_named_items as _build_named_items_util,
    first_signing as _first_signing_util,
    is_compaction_like as _is_compaction_like_util,
    is_flatness_like as _is_flatness_like_util,
    normalize_report_type as _normalize_report_type_util,
    normalize_standard_op as _normalize_standard_op_util,
    pick_template_name as _pick_template_name_util,
    resolve_schema_mode as _resolve_schema_mode_util,
    resolve_standard_rule as _resolve_standard_rule_util,
    resolve_table_headers as _resolve_table_headers_util,
    values_single_line as _values_single_line_util,
)
from services.api.docx_engine_signature_utils import (
    insert_signature as _insert_signature_util,
)
from services.api.docx_engine_utils import (
    coerce_values as _coerce_values_util,
    extract_unit as _extract_unit_util,
    format_display_time as _format_display_time_util,
    format_num as _format_num_util,
    format_signed_at as _format_signed_at_util,
    format_values_multiline as _format_values_multiline_util,
    normalize_payload as _normalize_payload_util,
    to_float as _to_float_util,
    to_text as _to_text_util,
    with_unit as _with_unit_util,
)

TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"
TEMPLATE_BY_TYPE = {
    "inspection": "01_inspection_report.docx",
    "lab": "02_lab_report.docx",
    "monthly_summary": "03_monthly_summary.docx",
    "final_archive": "04_final_archive_cover.docx",
}

PASS_CN = "\u5408\u683c"
FAIL_CN = "\u4e0d\u5408\u683c"
OBSERVE_CN = "\u89c2\u5bdf"
PENDING_CN = "\u5f85\u5b9a"
CANCELLED_CN = "\u53d6\u6d88"
PENDING_ANCHOR_CN = "\u5f85\u951a\u5b9a"
REBAR_SPACING_CN = "\u94a2\u7b4b\u95f4\u8ddd"
ROAD_FLATNESS_CN = "\u8def\u9762\u5e73\u6574\u5ea6"
SEPARATOR_CN = "\u3001"

SCHEMA_MODE_DESIGN_LIMIT = "design_limit"
SCHEMA_MODE_VALUE_STANDARD_MIN = "value_standard_min"
STANDARD_OP_PLUS_MINUS = "±"
FAIL_RGB = RGBColor(0xDC, 0x26, 0x26)


class DocxEngine:
    def _load(self, name: str) -> DocxTemplate:
        return DocxTemplate(str(TEMPLATE_DIR / name))

    def _qr(self, tpl: DocxTemplate, uri: str, size_mm: int = 22) -> InlineImage:
        """Generate QR code inline image for docxtpl."""
        img = qrcode.make(uri)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return InlineImage(tpl, buf, width=Mm(size_mm))

    def _fmt_result(self, result: str) -> str:
        return {
            "PASS": PASS_CN,
            "FAIL": FAIL_CN,
            "OBSERVE": OBSERVE_CN,
            "PENDING": PENDING_CN,
        }.get(result, result)

    def _now(self) -> str:
        return datetime.now().strftime("%Y-%m-%d %H:%M")

    def _now_seconds(self) -> str:
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def _verify_base_url(self) -> str:
        base = self._to_text(os.getenv("QCSPEC_VERIFY_BASE_URL") or "https://verify.qcspec.com").strip()
        return base.rstrip("/")

    @staticmethod
    def _to_text(value: Any, default: str = "") -> str:
        return _to_text_util(value, default)

    def _normalize_payload(self, value: Any) -> Any:
        return _normalize_payload_util(value)

    @staticmethod
    def _to_float(value: Any) -> float | None:
        return _to_float_util(value)

    @staticmethod
    def _coerce_values(values_raw: Any, fallback_value: Any = None) -> list[float]:
        return _coerce_values_util(values_raw, fallback_value=fallback_value)

    def _extract_unit(self, state_data: dict[str, Any]) -> str:
        return _extract_unit_util(state_data)

    def _with_unit(self, value_text: str, unit: str, *, force_inline: bool = False) -> str:
        return _with_unit_util(value_text, unit, force_inline=force_inline)

    @staticmethod
    def _format_num(value: float) -> str:
        return _format_num_util(value)

    def _format_values_multiline(self, values: list[float], chunk: int = 10, unit: str = "", force_inline_unit: bool = False) -> str | Listing:
        return _format_values_multiline_util(
            values,
            chunk=chunk,
            unit=unit,
            force_inline_unit=force_inline_unit,
            separator=SEPARATOR_CN,
        )

    def _resolve_test_type(self, state_data: dict[str, Any], *, fallback_type: str) -> tuple[str, str]:
        return _resolve_test_type_util(
            self,
            state_data,
            fallback_type=fallback_type,
            road_flatness_cn=ROAD_FLATNESS_CN,
            rebar_spacing_cn=REBAR_SPACING_CN,
        )

    def _extract_executor_name(self, signing: dict[str, Any], *, fallback_uri: str) -> str:
        return _extract_executor_name_util(self, signing, fallback_uri=fallback_uri)

    def _extract_executor_id(self, signing: dict[str, Any], *, fallback_uri: str, fallback_name: str) -> str:
        return _extract_executor_id_util(
            self,
            signing,
            fallback_uri=fallback_uri,
            fallback_name=fallback_name,
        )

    def _insert_signature(
        self,
        executor_id: str,
        tpl: DocxTemplate | None = None,
        size_mm: int = 18,
    ) -> tuple[InlineImage | bytes | str, str]:
        return _insert_signature_util(executor_id, tpl=tpl, size_mm=size_mm)

    def _build_v_uri_tree(
        self,
        *,
        project_uri: str,
        segment_uri: str,
        v_uri: str,
        proof_id: str,
        stake: str,
        verify_uri: str,
    ) -> dict[str, Any]:
        return _build_v_uri_tree_util(
            self,
            project_uri=project_uri,
            segment_uri=segment_uri,
            v_uri=v_uri,
            proof_id=proof_id,
            stake=stake,
            verify_uri=verify_uri,
        )

    def _resolve_schema_mode(self, state_data: dict[str, Any], *, test_type: str, test_type_name: str) -> str:
        return _resolve_schema_mode_util(
            state_data,
            test_type=test_type,
            test_type_name=test_type_name,
        )

    def _format_display_time(self, value: Any) -> str:
        return _format_display_time_util(value)

    def _format_signed_at(self, value: Any) -> str:
        return _format_signed_at_util(value)

    @staticmethod
    def _is_flatness_like(state_data: dict[str, Any]) -> bool:
        return _is_flatness_like_util(state_data)

    @staticmethod
    def _is_compaction_like(state_data: dict[str, Any]) -> bool:
        return _is_compaction_like_util(state_data)

    def _normalize_standard_op(self, raw_op: Any) -> str:
        return _normalize_standard_op_util(raw_op, plus_minus=STANDARD_OP_PLUS_MINUS)

    def _resolve_standard_rule(
        self,
        *,
        state_data: dict[str, Any],
        test_type: str,
        test_type_name: str,
        schema_mode: str,
        design: float | None,
        standard: float | None,
    ) -> tuple[str, float | None, float | None]:
        return _resolve_standard_rule_util(
            state_data=state_data,
            test_type=test_type,
            test_type_name=test_type_name,
            schema_mode=schema_mode,
            design=design,
            standard=standard,
            plus_minus=STANDARD_OP_PLUS_MINUS,
        )

    @staticmethod
    def _first_signing(proof: dict[str, Any]) -> dict[str, Any]:
        return _first_signing_util(proof)

    @staticmethod
    def _normalize_report_type(report_type: Any) -> str:
        return _normalize_report_type_util(report_type, template_by_type=TEMPLATE_BY_TYPE)

    def _pick_template_name(self, project_meta: dict[str, Any], report_type: str) -> str:
        return _pick_template_name_util(
            project_meta,
            report_type=report_type,
            template_by_type=TEMPLATE_BY_TYPE,
        )

    def _result_cn(self, result_code: str) -> str:
        code = self._to_text(result_code).strip().upper()
        if code == "CANCELLED":
            return CANCELLED_CN
        return self._fmt_result(code)

    def _build_named_items(self, rows: list[dict[str, Any]]) -> dict[str, Any]:
        return _build_named_items_util(rows)

    def _normalize_inspection_proof(self, proof: dict[str, Any]) -> dict[str, Any]:
        return _normalize_inspection_proof_util(
            self,
            proof,
            standard_op_plus_minus=STANDARD_OP_PLUS_MINUS,
            schema_mode_value_standard_min=SCHEMA_MODE_VALUE_STANDARD_MIN,
        )

    def _values_single_line(self, row: dict[str, Any]) -> str:
        return _values_single_line_util(row, separator=SEPARATOR_CN)

    def _resolve_table_headers(self, project_meta: dict[str, Any]) -> list[str]:
        return _resolve_table_headers_util(project_meta)

    @staticmethod
    def _set_cell_text(cell: Any, text: str, *, color: RGBColor | None = None) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(text or "-")
        if color is not None:
            run.font.color.rgb = color

    def _rewrite_inspection_table_rows(
        self,
        doc_bytes: bytes,
        rows: list[dict[str, Any]],
        *,
        header_labels: list[str] | None = None,
    ) -> bytes:
        return _rewrite_inspection_table_rows_util(
            self,
            doc_bytes,
            rows,
            header_labels=header_labels,
            fail_cn=FAIL_CN,
            fail_rgb=FAIL_RGB,
        )

    def render_docpeg_report(
        self,
        proofs: list[dict[str, Any]],
        project_meta: dict[str, Any],
        *,
        report_type: str = "inspection",
    ) -> bytes:
        """
        Generic sovereign export engine for QCSpec report templates.
        """
        project_meta = self._normalize_payload(project_meta)
        proofs = self._normalize_payload(proofs)

        normalized_type = self._normalize_report_type(report_type)
        template_name = self._pick_template_name(project_meta, normalized_type)
        tpl = self._load(template_name)

        rows: list[dict[str, Any]] = []
        any_fail = False
        sorted_proofs = sorted(
            [p for p in proofs if isinstance(p, dict)],
            key=lambda item: self._to_text(item.get("created_at")),
        )

        for idx, proof in enumerate(sorted_proofs, start=1):
            row, has_fail = _build_docpeg_row_util(
                self,
                proof,
                project_meta,
                idx=idx,
                normalized_type=normalized_type,
                schema_mode_design_limit=SCHEMA_MODE_DESIGN_LIMIT,
                standard_op_plus_minus=STANDARD_OP_PLUS_MINUS,
                pending_anchor_cn=PENDING_ANCHOR_CN,
            )
            rows.append(row)
            any_fail = any_fail or has_fail

        context = _build_docpeg_context_util(
            self,
            rows,
            project_meta,
            normalized_type=normalized_type,
            any_fail=any_fail,
            fail_cn=FAIL_CN,
            pass_cn=PASS_CN,
            pending_anchor_cn=PENDING_ANCHOR_CN,
        )
        verify_uri = self._to_text(context.get("verify_uri") or "")
        if verify_uri:
            qr = self._qr(tpl, verify_uri, size_mm=25)
            context["verify_qr"] = qr
            context["qr_image"] = qr
        signature_image, signature_status = self._insert_signature(
            self._to_text(context.get("executor_id") or context.get("signed_by") or ""),
            tpl=tpl,
        )
        context["signature_image"] = signature_image
        context["signature_loaded"] = signature_status == "loaded"
        context["signature_status"] = signature_status

        # Keep plain string substitution so rendered text inherits template run font.
        tpl.render(context, autoescape=False)
        buf = io.BytesIO()
        tpl.save(buf)
        rendered = buf.getvalue()
        if normalized_type == "inspection":
            return self._rewrite_inspection_table_rows(
                rendered,
                rows,
                header_labels=self._resolve_table_headers(project_meta),
            )
        return rendered

    def render_rebar_live_report(self, proofs: list, project_meta: dict) -> bytes:
        """
        Backward-compatible rebar report entry point.
        """
        return self.render_inspection_report(proofs, project_meta)

    def render_mode_driven_inspection_report(
        self,
        proofs: list[dict[str, Any]],
        project_meta: dict[str, Any],
    ) -> bytes:
        """
        Mode-driven generic inspection renderer.
        Supported schema modes:
        - design_limit: design + limit (+/-) + values
        - value_standard_max: value <= standard
        - value_standard_min: value >= standard
        - value_standard_eq: value == standard
        """
        normalized = [self._normalize_inspection_proof(p) for p in (proofs or []) if isinstance(p, dict)]
        return self.render_docpeg_report(normalized, project_meta, report_type="inspection")

    def render_universal_report(
        self,
        proofs: list[dict[str, Any]],
        project_meta: dict[str, Any],
        *,
        report_type: str = "inspection",
    ) -> bytes:
        """
        Universal schema-driven renderer for DocPeg.
        - Dynamically adapts test_type/value/unit/standard from proof.state_data
        - Produces atomic sovereignty variables for Word template binding
        - Keeps timestamp display at second precision (YYYY-MM-DD HH:mm:ss)
        """
        normalized_type = self._normalize_report_type(report_type)
        if normalized_type == "inspection":
            return self.render_mode_driven_inspection_report(proofs, project_meta)
        return self.render_docpeg_report(proofs, project_meta, report_type=normalized_type)

    def render_inspection_report(self, proofs: list[dict[str, Any]], project_meta: dict[str, Any]) -> bytes:
        """
        Explicit entrypoint for inspection report rendering.
        Handles dynamic inspection schema fields such as:
        - rebar_spacing (design/limit/values)
        - flatness-like tests (value/standard/unit)
        """
        return self.render_universal_report(proofs, project_meta, report_type="inspection")

